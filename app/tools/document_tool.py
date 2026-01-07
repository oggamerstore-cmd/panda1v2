"""
PANDA.1 Document Tool
=====================
Safe file browsing and preview for documents.

Version: 2.0

Supported formats:
- .docx (Word) via python-docx
- .txt (plain text)
- .md (Markdown)

Allowed roots:
- ~/Documents
- ~/.panda1/files
"""

import os
import logging
from pathlib import Path
from typing import Optional, List, Dict, Any
from dataclasses import dataclass
from html import escape

logger = logging.getLogger(__name__)

# Try to import python-docx
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.info("python-docx not installed, .docx preview disabled")


@dataclass
class DocumentResult:
    """Result of document operation."""
    success: bool
    path: Optional[str] = None
    name: Optional[str] = None
    content: Optional[str] = None
    html: Optional[str] = None
    summary: Optional[str] = None
    word_count: int = 0
    error: Optional[str] = None


class DocumentTool:
    """
    Safe document browsing and preview tool.
    
    Features:
    - Safe path resolution (no directory traversal)
    - Multiple format support
    - HTML preview generation
    - Content extraction for LLM processing
    """
    
    # Allowed root directories
    ALLOWED_ROOTS = [
        Path.home() / "Documents",
        Path.home() / ".panda1" / "files",
    ]
    
    # Supported extensions
    SUPPORTED_EXTENSIONS = {".docx", ".txt", ".md", ".text"}
    
    def __init__(self, additional_roots: Optional[List[Path]] = None):
        """
        Initialize document tool.
        
        Args:
            additional_roots: Additional allowed root directories
        """
        self.roots = list(self.ALLOWED_ROOTS)
        if additional_roots:
            self.roots.extend(additional_roots)
        
        # Ensure roots exist
        for root in self.roots:
            root.mkdir(parents=True, exist_ok=True)
    
    def _is_path_safe(self, path: Path) -> bool:
        """Check if path is within allowed roots."""
        try:
            resolved = path.resolve()
            for root in self.roots:
                try:
                    resolved.relative_to(root.resolve())
                    return True
                except ValueError:
                    continue
            return False
        except Exception:
            return False
    
    def _is_supported(self, path: Path) -> bool:
        """Check if file extension is supported."""
        return path.suffix.lower() in self.SUPPORTED_EXTENSIONS
    
    def list_files(self, root_index: int = 0) -> List[Dict[str, Any]]:
        """
        List files in an allowed root.
        
        Args:
            root_index: Index of root directory (0=Documents, 1=.panda1/files)
        
        Returns:
            List of file info dicts
        """
        if root_index >= len(self.roots):
            return []
        
        root = self.roots[root_index]
        files = []
        
        try:
            for item in root.rglob("*"):
                if item.is_file() and self._is_supported(item):
                    try:
                        stat = item.stat()
                        files.append({
                            "name": item.name,
                            "path": str(item),
                            "relative_path": str(item.relative_to(root)),
                            "size": stat.st_size,
                            "modified": stat.st_mtime,
                            "extension": item.suffix.lower(),
                        })
                    except Exception as e:
                        logger.debug(f"Error getting file info: {e}")
        except Exception as e:
            logger.error(f"Error listing files: {e}")
        
        # Sort by modified time (newest first)
        files.sort(key=lambda x: x["modified"], reverse=True)
        
        return files
    
    def open_file(self, path: str) -> DocumentResult:
        """
        Open and read a document.
        
        Args:
            path: Path to document
        
        Returns:
            DocumentResult with content
        """
        file_path = Path(path)
        
        # Security check
        if not self._is_path_safe(file_path):
            return DocumentResult(
                success=False,
                error="Access denied: path outside allowed directories"
            )
        
        if not file_path.exists():
            return DocumentResult(
                success=False,
                error=f"File not found: {file_path.name}"
            )
        
        if not self._is_supported(file_path):
            return DocumentResult(
                success=False,
                error=f"Unsupported file type: {file_path.suffix}"
            )
        
        ext = file_path.suffix.lower()
        
        try:
            if ext == ".docx":
                return self._read_docx(file_path)
            elif ext in (".txt", ".text", ".md"):
                return self._read_text(file_path)
            else:
                return DocumentResult(
                    success=False,
                    error=f"Unsupported extension: {ext}"
                )
        except Exception as e:
            logger.error(f"Error reading file: {e}")
            return DocumentResult(
                success=False,
                error=str(e)
            )
    
    def _read_docx(self, path: Path) -> DocumentResult:
        """Read a Word document."""
        if not DOCX_AVAILABLE:
            return DocumentResult(
                success=False,
                error="python-docx not installed. Run: pip install python-docx"
            )
        
        doc = DocxDocument(str(path))
        
        # Extract text
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        
        content = "\n\n".join(paragraphs)
        word_count = len(content.split())
        
        # Generate HTML preview
        html_parts = ['<div class="docx-preview">']
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                style = para.style.name if para.style else ""
                if "Heading" in style:
                    level = style.replace("Heading", "").strip() or "2"
                    try:
                        level = min(6, max(1, int(level)))
                    except ValueError:
                        level = 2
                    html_parts.append(f"<h{level}>{escape(text)}</h{level}>")
                else:
                    html_parts.append(f"<p>{escape(text)}</p>")
        html_parts.append('</div>')
        html = "\n".join(html_parts)
        
        return DocumentResult(
            success=True,
            path=str(path),
            name=path.name,
            content=content,
            html=html,
            word_count=word_count,
        )
    
    def _read_text(self, path: Path) -> DocumentResult:
        """Read a text file."""
        # Try common encodings
        encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252"]
        
        content = None
        for encoding in encodings:
            try:
                with open(path, "r", encoding=encoding) as f:
                    content = f.read()
                break
            except UnicodeDecodeError:
                continue
        
        if content is None:
            return DocumentResult(
                success=False,
                error="Unable to decode file (tried UTF-8, Latin-1, CP1252)"
            )
        
        word_count = len(content.split())
        
        # Generate HTML
        if path.suffix.lower() == ".md":
            # Basic markdown-ish rendering
            html = self._render_markdown(content)
        else:
            # Plain text
            html = f'<pre class="text-preview">{escape(content)}</pre>'
        
        return DocumentResult(
            success=True,
            path=str(path),
            name=path.name,
            content=content,
            html=html,
            word_count=word_count,
        )
    
    def _render_markdown(self, content: str) -> str:
        """Basic markdown to HTML rendering."""
        lines = content.split("\n")
        html_parts = ['<div class="md-preview">']
        
        in_code = False
        code_lines = []
        
        for line in lines:
            # Code blocks
            if line.startswith("```"):
                if in_code:
                    html_parts.append(f'<pre><code>{escape(chr(10).join(code_lines))}</code></pre>')
                    code_lines = []
                    in_code = False
                else:
                    in_code = True
                continue
            
            if in_code:
                code_lines.append(line)
                continue
            
            # Headers
            if line.startswith("# "):
                html_parts.append(f"<h1>{escape(line[2:])}</h1>")
            elif line.startswith("## "):
                html_parts.append(f"<h2>{escape(line[3:])}</h2>")
            elif line.startswith("### "):
                html_parts.append(f"<h3>{escape(line[4:])}</h3>")
            elif line.strip():
                html_parts.append(f"<p>{escape(line)}</p>")
        
        html_parts.append('</div>')
        return "\n".join(html_parts)
    
    def summarize(self, path: str, max_words: int = 200) -> DocumentResult:
        """
        Get a summary of document (first N words).
        
        Args:
            path: Path to document
            max_words: Maximum words to extract
        
        Returns:
            DocumentResult with summary
        """
        result = self.open_file(path)
        
        if not result.success:
            return result
        
        # Extract first N words
        words = result.content.split()
        if len(words) > max_words:
            summary = " ".join(words[:max_words]) + "..."
        else:
            summary = result.content
        
        result.summary = summary
        return result


# Global tool instance
_doc_tool: Optional[DocumentTool] = None


def get_document_tool() -> DocumentTool:
    """Get or create the global document tool."""
    global _doc_tool
    if _doc_tool is None:
        _doc_tool = DocumentTool()
    return _doc_tool


def list_documents(root_index: int = 0) -> List[Dict[str, Any]]:
    """List available documents."""
    return get_document_tool().list_files(root_index)


def open_document(path: str) -> DocumentResult:
    """Open a document."""
    return get_document_tool().open_file(path)


def summarize_document(path: str, max_words: int = 200) -> DocumentResult:
    """Get document summary."""
    return get_document_tool().summarize(path, max_words)
